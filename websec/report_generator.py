from __future__ import annotations

import copy
import datetime as dt
import io
import os
import re
import shutil
import ssl
import subprocess
import tempfile
import urllib.parse
import urllib.request
import zipfile
from pathlib import Path
from typing import Callable, Dict, Iterable, List, Optional

from docx import Document
from docx.text.paragraph import Paragraph
from PIL import Image

from owasp_mapper import SECTIONS, group_findings_by_section, severity_counts
from screenshot import capture_homepage_screenshot
from whois_parser import DomainWhoisInfo, lookup_domain_info


DEFAULT_TEMPLATE_CANDIDATES = [
    os.environ.get("DEEPSEC_REPORT_TEMPLATE", "").strip(),
    r"E:\default_report\default_report.docx",
    str(Path(__file__).resolve().parent / "report_settings" / "default_report.docx"),
    str(Path(__file__).resolve().parent / "default_report.docx"),
]

DEFAULT_COVER_CANDIDATES = [
    os.environ.get("DEEPSEC_REPORT_DEFAULT_COVER", "").strip(),
    str(Path(__file__).resolve().parent / "report_settings" / "defualt_page.png"),
    str(Path(__file__).resolve().parent / "report_settings" / "default_page.png"),
    str(Path(__file__).resolve().parent / "defualt_page.png"),
    str(Path(__file__).resolve().parent / "default_page.png"),
]

SECTION_PARAGRAPH_INDEXES = {
    "cover_site": 6,
    "intro_body": 60,
    "intro_caption": 63,
    "domain_heading": 78,
    "results_heading": 80,
    "sec61_placeholder": 91,
    "sec62_detail_1": 96,
    "sec62_detail_2": 97,
    "sec62_detail_3": 98,
    "sec62_detail_4": 99,
    "sec62_caption_2": 102,
    "sec62_caption_3": 104,
    "sec62_caption_4": 106,
    "sec62_caption_5": 108,
    "sec62_recommendation": 111,
    "sec62_password_intro": 112,
    "sec62_password_1": 113,
    "sec62_password_2": 114,
    "sec62_password_3": 115,
    "sec62_password_4": 116,
    "sec63_placeholder": 121,
    "sec64_placeholder": 126,
    "sec65_placeholder": 131,
    "sec66_placeholder": 136,
    "sec67_placeholder": 141,
    "sec68_placeholder": 145,
    "sec69_placeholder": 151,
    "sec610_placeholder": 157,
}

SECTION_PLACEHOLDER_MAP = {
    "6.1": "sec61_placeholder",
    "6.3": "sec63_placeholder",
    "6.4": "sec64_placeholder",
    "6.5": "sec65_placeholder",
    "6.6": "sec66_placeholder",
    "6.7": "sec67_placeholder",
    "6.8": "sec68_placeholder",
    "6.9": "sec69_placeholder",
    "6.10": "sec610_placeholder",
}

MEDIA_REPLACEMENTS = {
    "cover_page": "word/media/image1.png",
    "intro_screenshot": "word/media/image2.png",
    "blank_slots": [
        "word/media/image3.png",
        "word/media/image4.png",
        "word/media/image5.png",
        "word/media/image6.jpeg",
        "word/media/image7.png",
    ],
}

UZ_MONTHS = {
    1: "yanvar",
    2: "fevral",
    3: "mart",
    4: "aprel",
    5: "may",
    6: "iyun",
    7: "iyul",
    8: "avgust",
    9: "sentyabr",
    10: "oktyabr",
    11: "noyabr",
    12: "dekabr",
}


def _log(logger: Optional[Callable[[str], None]], message: str) -> None:
    if callable(logger):
        logger(message)


def _attr(item, name: str, default=""):
    if isinstance(item, dict):
        return item.get(name, default)
    return getattr(item, name, default)


def resolve_template_path(template_path: str | Path | None = None) -> Optional[Path]:
    candidates = []
    if template_path:
        candidates.append(str(template_path))
    candidates.extend(DEFAULT_TEMPLATE_CANDIDATES)
    for raw in candidates:
        value = str(raw or "").strip()
        if not value:
            continue
        path = Path(value)
        if path.exists() and path.is_file():
            return path
    return None


def resolve_cover_image_path(cover_path: str | Path | None = None) -> Optional[Path]:
    candidates = []
    if cover_path:
        candidates.append(str(cover_path))
    candidates.extend(DEFAULT_COVER_CANDIDATES)
    for raw in candidates:
        value = str(raw or "").strip()
        if not value:
            continue
        path = Path(value)
        if path.exists() and path.is_file():
            return path
    return None


def _parse_datetime(value) -> Optional[dt.datetime]:
    if isinstance(value, dt.datetime):
        return value
    text = str(value or "").strip()
    if not text:
        return None
    text = text.replace("Z", "+00:00")
    try:
        return dt.datetime.fromisoformat(text)
    except Exception:
        return None


def _format_short_date(value: dt.datetime) -> str:
    return value.strftime("%d.%m.%Y")


def _format_long_uzbek_date(value: dt.datetime) -> str:
    return f"{value.year}-yil {value.day}-{UZ_MONTHS.get(value.month, value.month)}"


def _business_days(start: dt.datetime, end: dt.datetime) -> int:
    begin = start.date()
    finish = end.date()
    if finish < begin:
        begin, finish = finish, begin
    current = begin
    count = 0
    while current <= finish:
        if current.weekday() < 5:
            count += 1
        current += dt.timedelta(days=1)
    return max(count, 1)


def _extract_target_parts(target: str) -> dict:
    normalized = str(target or "").strip()
    if not urllib.parse.urlparse(normalized).scheme:
        normalized = f"https://{normalized}"
    parsed = urllib.parse.urlparse(normalized)
    host = (parsed.netloc or parsed.path).strip()
    base_domain = host[4:] if host.startswith("www.") else host
    www_domain = f"www.{base_domain}" if base_domain and "." in base_domain and not host.startswith("www.") else host
    return {
        "url": normalized,
        "host": host or base_domain,
        "base_domain": base_domain or host,
        "www_domain": www_domain or host or base_domain,
    }


def _safe_name(value: str) -> str:
    return re.sub(r"[^\w.-]+", "_", str(value or "").strip()) or "report"


def _probe_reachable_url(url: str) -> Optional[str]:
    normalized = str(url or "").strip()
    if not normalized:
        return None
    request = urllib.request.Request(
        normalized,
        headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/135.0.0.0 Safari/537.36"
        },
    )
    context = ssl.create_default_context()
    context.check_hostname = False
    context.verify_mode = ssl.CERT_NONE
    opener = urllib.request.build_opener(urllib.request.HTTPSHandler(context=context))
    try:
        with opener.open(request, timeout=25) as response:
            final_url = str(response.geturl() or normalized)
            status = int(getattr(response, "status", 200) or 200)
            content_type = str(response.headers.get("Content-Type", "")).lower()
            if status >= 400:
                return None
            if "text/html" not in content_type and "application/xhtml+xml" not in content_type:
                return None
            return final_url
    except Exception:
        return None


def _resolve_screenshot_url(target_parts: dict, logger: Optional[Callable[[str], None]] = None) -> str:
    candidates: List[str] = []
    original = str(target_parts.get("url") or "").strip()
    base_domain = str(target_parts.get("base_domain") or "").strip()
    www_domain = str(target_parts.get("www_domain") or "").strip()

    for candidate in [
        original,
        f"https://{base_domain}" if base_domain else "",
        f"https://{www_domain}" if www_domain else "",
        f"http://{base_domain}" if base_domain else "",
        f"http://{www_domain}" if www_domain else "",
    ]:
        value = str(candidate or "").strip()
        if value and value not in candidates:
            candidates.append(value)

    for candidate in candidates:
        resolved = _probe_reachable_url(candidate)
        if resolved:
            if resolved != original:
                _log(logger, f"Report: screenshot source switched from {original} to {resolved}")
            return resolved

    return original


def _friendly_tech_value(value: str, *, upper: bool = False) -> str:
    text = str(value or "").strip()
    if not text or text.lower() == "unknown":
        return "Aniqlanmadi"
    if upper:
        return text.upper()
    mapping = {
        "php": "PHP",
        "python": "Python",
        "javascript": "JavaScript",
        "node": "Node.js",
        "nginx": "NGINX",
        "apache": "Apache",
        "iis": "IIS",
        "yii2": "Yii2",
        "codeigniter": "CodeIgniter",
        "laravel": "Laravel",
        "wordpress": "WordPress",
        "django": "Django",
        "flask": "Flask",
        "fastapi": "FastAPI",
        "spring": "Spring",
    }
    return mapping.get(text.lower(), text)


def _build_anchor_map(doc: Document) -> Dict[str, Paragraph]:
    return {key: doc.paragraphs[index] for key, index in SECTION_PARAGRAPH_INDEXES.items()}


def _anchor(context: dict, key: str) -> Paragraph:
    return context["anchors"][key]


def _run_has_drawing(run) -> bool:
    return bool(run._element.xpath(".//*[local-name()='drawing']"))


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    editable_runs = [run for run in paragraph.runs if not _run_has_drawing(run)]
    if editable_runs:
        editable_runs[0].text = text
        for run in editable_runs[1:]:
            run.text = ""
    elif paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = None
    for candidate in paragraph._parent.paragraphs:
        if candidate._p is new_p:
            new_para = candidate
            break
    if new_para is None:
        new_para = paragraph._parent.paragraphs[-1]
    _set_paragraph_text(new_para, text)
    new_para.style = paragraph.style
    return new_para


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _update_table_cell(cell, value: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph()
    first = cell.paragraphs[0]
    _set_paragraph_text(first, value)
    for extra in list(cell.paragraphs[1:]):
        _delete_paragraph(extra)


def _table_row_by_label(table, label: str):
    def _norm(value: str) -> str:
        text = str(value or "").lower()
        for src in ("’", "‘", "ʻ", "ʼ", "`"):
            text = text.replace(src, "'")
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    expected = _norm(label)
    for row in table.rows:
        if row.cells and _norm(row.cells[0].text) == expected:
            return row
    return None


def _clear_rows_except_header(table) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)


def _risk_uz_label(risk: str) -> str:
    normalized = str(risk or "").strip().lower()
    if normalized in {"critical", "high"}:
        return "Yuqori"
    if normalized == "medium":
        return "O'rta"
    return "Quyi"


def _shorten(text: str, limit: int = 420) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "").strip())
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[: limit - 3].rstrip() + "..."


def _clean_report_text(text: str) -> str:
    value = str(text or "")
    replacements = {
        "вЂ™": "'",
        "вЂ": "'",
        "вЂњ": '"',
        "вЂќ": '"',
        "maвЂ™lumot": "ma'lumot",
        "boвЂyicha": "bo'yicha",
        "oвЂtkazildi": "o'tkazildi",
        "oвЂtkazilgan": "o'tkazilgan",
        "koвЂra": "ko'ra",
        "taКј": "ta'",
        "OК»": "O'",
        "К»": "'",
        "Кј": "'",
    }
    for src, dst in replacements.items():
        value = value.replace(src, dst)
    return re.sub(r"\s+", " ", value).strip()


def _finding_haystack(finding: object) -> str:
    return " ".join(
        _clean_report_text(_attr(finding, field, "")).lower()
        for field in ("owasp_id", "owasp_name", "title", "evidence", "tool", "payload", "url")
    )


def _matches_finding(finding: object, *keywords: str) -> bool:
    haystack = _finding_haystack(finding)
    return any(keyword.lower() in haystack for keyword in keywords if keyword)


def _localized_title(finding: object) -> str:
    raw_title = _clean_report_text(_attr(finding, "title", ""))
    title_l = raw_title.lower()

    if _matches_finding(finding, "broken access control", "idor", "403 bypass", "bac", "forbidden parent"):
        if "configure" in title_l or "config" in title_l:
            return "Himoyalangan konfiguratsiya sahifasiga ruxsatsiz kirish"
        if "ip bypass" in title_l:
            return "IP spoofing header orqali himoyani chetlab o'tish"
        if "url override" in title_l:
            return "URL override header orqali himoyani chetlab o'tish"
        if "path variant" in title_l:
            return "Path varianti orqali himoyani chetlab o'tish"
        return "Foydalana olish nazoratining buzilishi"

    if _matches_finding(finding, "sqli", "sql injection", "sqlmap"):
        return "SQL injeksiya alomati"
    if _matches_finding(finding, "stored xss"):
        return "Stored XSS alomati"
    if _matches_finding(finding, "xss", "dalfox"):
        return "XSS alomati"
    if _matches_finding(finding, "ssrf", "oob"):
        return "SSRF alomati"
    if _matches_finding(finding, "lfi", "path traversal", "etc/passwd", "win.ini"):
        return "Fayl o'qish yoki path traversal alomati"
    if _matches_finding(finding, "cmdi", "command injection", "commix", "rce"):
        return "Buyruq injeksiyasi alomati"
    if _matches_finding(finding, "ssti", "template injection", "jinja", "twig", "freemarker"):
        return "Server-side template injection alomati"
    if _matches_finding(finding, "jwt", "alg:none", "hs256"):
        return "JWT xavfsizligi bilan bog'liq zaiflik"
    if _matches_finding(finding, "graphql"):
        return "GraphQL bilan bog'liq zaiflik"
    if _matches_finding(finding, "csrf"):
        return "CSRF himoyasi yetishmasligi"
    if _matches_finding(finding, "default credential", "default creds"):
        return "Standart credentiallar bilan kirish"
    if _matches_finding(finding, "rate limit", "brute force"):
        return "Rate limiting yetishmasligi"
    if _matches_finding(finding, "security misconfiguration", "misconfig", "phpmyadmin", "adminer", "debug", "exposed", "directory listing"):
        return "Noto'g'ri xavfsizlik sozlamasi"
    if _matches_finding(finding, "tls", "ssl", "weak cipher", "certificate"):
        return "Kriptografik himoya zaifligi"
    if _matches_finding(finding, "cookie", "session fixation", "samesite", "httponly"):
        return "Sessiya yoki cookie xavfsizligi zaifligi"

    if raw_title:
        basic = raw_title
        phrase_map = [
            ("Reason:", "Sabab:"),
            ("Evidence:", "Dalil:"),
            ("Forbidden parent, child accessible", "Yopiq ota manzil ostidagi child resurs ochiq"),
            ("Possible IDOR", "IDOR ehtimoli"),
            ("Security Misconfiguration", "Noto'g'ri xavfsizlik sozlamasi"),
        ]
        for src, dst in phrase_map:
            basic = basic.replace(src, dst)
        return basic
    return "Zaiflik"


def _localize_raw_fragment(text: str) -> str:
    localized = _clean_report_text(text)
    replacements = [
        ("Reason:", "Sabab:"),
        ("Evidence:", "Dalil:"),
        ("real bac", "haqiqiy BAC holati"),
        ("body size", "javob hajmi"),
        ("sensitive data", "sezgir ma'lumotlar"),
        ("credentials", "credentiallar"),
        ("secrets", "sirli qiymatlar"),
        ("system paths", "tizim yo'llari"),
        ("returns 200", "200 holat kodi bilan javob qaytardi"),
        ("returned 200", "200 holat kodi bilan javob qaytardi"),
        ("actual administrative configuration data", "haqiqiy administrator konfiguratsiyasi ma'lumotlari"),
        ("Different data returned", "Turli ma'lumot qaytdi"),
        ("time-based", "vaqtga asoslangan"),
        ("delay", "kechikish"),
        ("reflected", "aks etdi"),
        ("batch queries", "batch so'rovlar"),
        ("introspection", "introspection"),
        ("stack trace", "stack trace"),
        ("debug console", "debug console"),
        ("rate limiting", "rate limiting"),
        ("default credentials", "standart credentiallar"),
    ]
    for src, dst in replacements:
        localized = localized.replace(src, dst)
    return localized


def _evidence_text_for_finding(finding: object) -> str:
    raw_evidence = _clean_report_text(_attr(finding, "evidence", ""))
    raw_url = _clean_report_text(_attr(finding, "url", "")) or "aniqlangan manzil"
    raw_param = _clean_report_text(_attr(finding, "param", ""))

    if _matches_finding(finding, "broken access control", "idor", "403 bypass", "bac", "forbidden parent"):
        detail = "Tekshiruvda himoyalangan resurs ruxsatsiz ochilgani yoki kutilgan cheklov chetlab o'tilgani kuzatildi."
        if "configure" in raw_url or "config" in raw_url or "configure" in raw_evidence.lower():
            detail = "Tekshiruvda himoyalangan konfiguratsiya sahifasi 403/401 o'rniga ochilib, administratorga oid ma'lumotlar qaytgani kuzatildi."
        elif re.search(r"\b403\b", raw_evidence) and re.search(r"\b200\b", raw_evidence):
            detail = "Tekshiruvda himoyalangan endpoint uchun 403 o'rniga 200 javob qaytishi kuzatildi."
        if raw_param.startswith("header:"):
            detail += f" Tekshiruv header orqali amalga oshirilgan: {raw_param.split(':', 1)[1]}."
        return _shorten(detail, 260)

    if _matches_finding(finding, "sqli", "sql injection", "sqlmap"):
        if "sleep" in raw_evidence.lower() or "waitfor" in raw_evidence.lower() or "delay" in raw_evidence.lower():
            return "Kiritilgan payloaddan so'ng vaqtga asoslangan anomaliya kuzatildi, bu backend SQL qayta ishlashiga ta'sir bo'lishi mumkinligini ko'rsatadi."
        if any(token in raw_evidence.lower() for token in ["sql", "mysql", "syntax", "ora-", "postgres", "sqlite"]):
            return f"Javobda SQL bilan bog'liq xatolik yoki aniq farq kuzatildi. Qo'shimcha signal: {_shorten(_localize_raw_fragment(raw_evidence), 160)}"
        return "Kiritilgan qiymat backend SQL qayta ishlash jarayoniga ta'sir qilishi mumkinligini ko'rsatuvchi farqlar aniqlandi."

    if _matches_finding(finding, "stored xss"):
        return "Yuborilgan XSS markeri keyingi sahifada saqlanib qayta aks etgani kuzatildi."

    if _matches_finding(finding, "xss", "dalfox"):
        return "Kiritilgan payload sahifa javobida yoki brauzerda bajariladigan kontekstda qayta aks etish alomatini ko'rsatdi."

    if _matches_finding(finding, "lfi", "path traversal", "etc/passwd", "win.ini"):
        if any(token in raw_evidence.lower() for token in ["root:", "/bin/", "[extensions]", "win.ini"]):
            return f"Javobda server fayliga oid mazmun qaytgani kuzatildi. Qo'shimcha signal: {_shorten(_localize_raw_fragment(raw_evidence), 150)}"
        return "Fayl yoki yo'l parametri orqali serverdagi lokal fayllarga murojaat qilish alomati kuzatildi."

    if _matches_finding(finding, "ssrf", "oob"):
        if "oob" in raw_evidence.lower():
            return "Tashqi OOB kuzatuv nuqtasiga callback kelgani sababli server tashqi yoki ichki manzilga so'rov yuborgan bo'lishi mumkin."
        if any(token in raw_evidence.lower() for token in ["169.254", "instance-id", "ami-id", "metadata"]):
            return "Javobda ichki servis yoki metadata endpointiga tegishli ma'lumotlar qaytgani kuzatildi."
        return "Server foydalanuvchi boshqaradigan URL manzilga so'rov yuborishi mumkinligini ko'rsatuvchi signal aniqlandi."

    if _matches_finding(finding, "cmdi", "command injection", "commix", "rce"):
        if any(token in raw_evidence.lower() for token in ["uid=", "gid=", "whoami", "/bin/", "root:"]):
            return f"Javobda tizim buyruqlariga xos natija ko'rindi. Qo'shimcha signal: {_shorten(_localize_raw_fragment(raw_evidence), 150)}"
        return "Parametr qiymati tizim buyruqlarini bajarish jarayoniga ta'sir qilishi mumkinligini ko'rsatuvchi alomat kuzatildi."

    if _matches_finding(finding, "ssti", "template injection", "jinja", "twig", "freemarker"):
        return "Template ifodasiga o'xshash payload kiritilganda server javobida hisoblangan yoki ishlov berilgan natija alomatlari kuzatildi."

    if _matches_finding(finding, "jwt", "alg:none", "hs256"):
        return "JWT tekshiruvida token algoritmi yoki imzo nazorati xavfsiz tarzda cheklanmagani bo'yicha signal aniqlandi."

    if _matches_finding(finding, "graphql"):
        if "introspection" in raw_evidence.lower():
            return "GraphQL introspection yoqilgan bo'lib, schema tuzilmasi tashqi foydalanuvchiga ko'rinmoqda."
        return "GraphQL endpointida xavfsizlikka ta'sir qiluvchi tekshiruv natijalari qayd etildi."

    if _matches_finding(finding, "csrf"):
        return "State-changing so'rov uchun CSRF himoya tokeni aniqlanmadi."

    if _matches_finding(finding, "default credential", "default creds"):
        return "Standart credentiallar yordamida tizimga kirish imkoni kuzatildi."

    if _matches_finding(finding, "rate limit", "brute force"):
        return "Qisqa vaqt ichida ko'p marta yuborilgan so'rovlar uchun yetarli cheklov yoki bloklash kuzatilmadi."

    if _matches_finding(finding, "security misconfiguration", "misconfig", "phpmyadmin", "adminer", "debug", "exposed", "directory listing"):
        return "Xavfsizlik sozlamalari noto'g'ri qo'llangani sababli xizmat, fayl yoki boshqaruv interfeysi ortiqcha ochiq qolgan."

    if _matches_finding(finding, "tls", "ssl", "weak cipher", "certificate"):
        return "TLS/SSL sozlamalarida eskirgan protokol yoki zaif shifrlash parametrlari ishlatilayotgani bo'yicha signal aniqlandi."

    if _matches_finding(finding, "cookie", "session fixation", "samesite", "httponly"):
        return "Sessiya yoki cookie bilan ishlashda himoya atributlari yetarli emasligi aniqlandi."

    if raw_evidence:
        return _shorten(_localize_raw_fragment(raw_evidence), 260)
    return "Dalil mavjud emas."


def _default_recommendation_for_finding(finding: object) -> str:
    if _matches_finding(finding, "broken access control", "idor", "403 bypass", "bac", "forbidden parent"):
        return (
            "Har bir endpoint, path varianti va HTTP method uchun server tomonda authorization tekshiruvini majburiy qo'llang. "
            "X-Forwarded-For, X-Original-URL va shunga o'xshash headerlarga ishonmang."
        )
    if _matches_finding(finding, "sqli", "sql injection", "sqlmap"):
        return (
            "SQL so'rovlarida parameterized query/prepared statementlardan foydalaning. "
            "Foydalanuvchi kiritgan qiymatlarni SQL satriga to'g'ridan-to'g'ri qo'shmang va kiruvchi ma'lumotni qat'iy validatsiya qiling."
        )
    if _matches_finding(finding, "stored xss", "xss", "dalfox"):
        return (
            "Foydalanuvchi kiritgan ma'lumotlarni chiqarishda contextga mos encode qiling, xavfli HTML/JS ni sanitizatsiya qiling "
            "va qo'shimcha himoya sifatida Content-Security-Policy joriy qiling."
        )
    if _matches_finding(finding, "lfi", "path traversal", "etc/passwd", "win.ini"):
        return (
            "Fayl yo'llarini allowlist asosida tekshiring, `../` kabi traversal belgilarini rad eting "
            "va konfiguratsiya hamda tizim fayllarini web root tashqarisida saqlang."
        )
    if _matches_finding(finding, "ssrf", "oob"):
        return (
            "Foydalanuvchi yuboradigan URL qiymatlarini qat'iy tekshiring, faqat ruxsat etilgan domen va protokollarga ruxsat bering "
            "hamda ichki IP diapazonlari va metadata endpointlariga chiqishni bloklang."
        )
    if _matches_finding(finding, "cmdi", "command injection", "commix", "rce"):
        return (
            "Foydalanuvchi kiritgan qiymatlarni shell buyruqlariga uzatmang. Tizim bilan ishlash uchun xavfsiz kutubxona/API ishlating "
            "va buyruq argumentlarini qat'iy allowlist asosida tekshiring."
        )
    if _matches_finding(finding, "ssti", "template injection", "jinja", "twig", "freemarker"):
        return (
            "Foydalanuvchi ma'lumotini template kodi sifatida qayta ishlamang, server-side rendering kontekstini cheklang "
            "va template engine xavfsiz rejimlarini yoqing."
        )
    if _matches_finding(finding, "jwt", "alg:none", "hs256"):
        return (
            "JWT uchun faqat ruxsat etilgan algoritmlarni whitelist qiling, `alg:none` ni rad eting "
            "va kuchli sir kalitlardan foydalanib tokenlarni rotatsiya qiling."
        )
    if _matches_finding(finding, "graphql"):
        return (
            "Production muhitida introspectionni o'chiring, query depth va complexity limitlarini joriy qiling "
            "hamda resolverlarda kiruvchi ma'lumotni alohida validatsiya qiling."
        )
    if _matches_finding(finding, "csrf"):
        return "Barcha state-changing forma va endpointlarga CSRF token qo'shing va server tomonda uni tekshiring."
    if _matches_finding(finding, "default credential", "default creds"):
        return "Standart credentiallarni darhol almashtiring, kuchli parol siyosatini joriy qiling va birinchi kirishda parolni yangilashni majburiy qiling."
    if _matches_finding(finding, "rate limit", "brute force"):
        return "Login va autentifikatsiya endpointlariga rate limiting, vaqtinchalik bloklash va kerak bo'lsa CAPTCHA himoyasini qo'llang."
    if _matches_finding(finding, "security misconfiguration", "misconfig", "phpmyadmin", "adminer", "debug", "exposed", "directory listing"):
        return (
            "Keraksiz servis, debug rejim, admin panel va ochiq fayllarni o'chiring yoki cheklang. "
            "Xizmatlarni faqat zarur tarmoq manzillari uchun ochib, xavfsiz default sozlamalarni qo'llang."
        )
    if _matches_finding(finding, "tls", "ssl", "weak cipher", "certificate"):
        return "TLS 1.2+ dan foydalaning, eskirgan protokollar va zaif cipher suite'larni o'chiring hamda sertifikat zanjirini to'g'ri sozlang."
    if _matches_finding(finding, "cookie", "session fixation", "samesite", "httponly"):
        return "Sessiya cookie'lariga HttpOnly, Secure va SameSite atributlarini qo'shing hamda autentifikatsiyadan keyin sessiya identifikatorini yangilang."
    return "Aniqlangan zaiflikni bartaraf etish uchun kiruvchi ma'lumotlarni tekshirish, server tomondagi nazoratlarni kuchaytirish va ruxsatlarni eng kam imtiyoz tamoyili asosida qayta ko'rib chiqish tavsiya etiladi."


def _localize_remediation_text(text: str, finding: object) -> str:
    raw = _clean_report_text(text)
    if not raw:
        return ""

    replacements = [
        ("Use parameterized queries/prepared statements.", "SQL so'rovlarida parameterized query/prepared statementlardan foydalaning."),
        ("Never concatenate user input into SQL.", "Foydalanuvchi kiritgan qiymatlarni SQL satriga to'g'ridan-to'g'ri qo'shmang."),
        ("Example (Python):", "Misol (Python):"),
        ("Additionally, implement input validation.", "Qo'shimcha ravishda kiruvchi ma'lumotlarni validatsiya qiling."),
        ("Validate and whitelist allowed URLs.", "URL qiymatlarini tekshiring va faqat ruxsat etilgan manzillar ro'yxatiga ruxsat bering."),
        ("Block internal IPs.", "Ichki IP diapazonlariga murojaatni bloklang."),
        ("Whitelist allowed domains.", "Faqat ruxsat etilgan domenlar ro'yxatiga ruxsat bering."),
        ("Validate URL scheme.", "URL sxemasini tekshiring."),
        ("Restrict access.", "Kirishni cheklang."),
        ("Disable introspection in production.", "Production muhitida introspectionni o'chiring."),
        ("Use query depth limiting.", "Query depth limiting joriy qiling."),
        ("Disable batch queries or add per-query rate limiting.", "Batch querylarni o'chiring yoki har bir so'rov uchun rate limiting joriy qiling."),
        ("Whitelist allowed redirect targets.", "Faqat ruxsat etilgan redirect manzillarini whitelist qiling."),
        ("Never use user input as redirect URL.", "Foydalanuvchi qiymatini to'g'ridan-to'g'ri redirect URL sifatida ishlatmang."),
        ("Validate redirect targets server-side.", "Redirect manzillarini server tomonda tekshiring."),
        ("Strip CR/LF characters from all user inputs used in HTTP headers.", "HTTP headerlarda ishlatiladigan barcha foydalanuvchi qiymatlaridan CR/LF belgilarini olib tashlang."),
        ("Add CSRF tokens to all state-changing forms.", "Barcha state-changing formalarga CSRF token qo'shing."),
        ("Implement rate limiting + CAPTCHA on auth endpoints.", "Auth endpointlarga rate limiting va kerak bo'lsa CAPTCHA himoyasini joriy qiling."),
        ("Whitelist allowed JWT algorithms. Reject alg:none.", "JWT uchun faqat ruxsat etilgan algoritmlarni whitelist qiling va `alg:none` ni rad eting."),
        ("Use asymmetric RS256/ES256.", "Imkon qadar asimmetrik RS256/ES256 algoritmlaridan foydalaning."),
        ("Disable external entity processing.", "Tashqi entity qayta ishlanishini o'chiring."),
        ("Move secrets to environment variables. Rotate compromised credentials immediately.", "Sir kalitlarni environment variable'larda saqlang va fosh bo'lgan credentiallarni zudlik bilan almashtiring."),
        ("Move to environment variables. Rotate if real.", "Bunday qiymatlarni environment variable'ga ko'chiring va haqiqiy bo'lsa darhol almashtiring."),
        ("Use signed JWT or encrypted tokens instead of serialized objects.", "Serialized obyektlar o'rniga imzolangan yoki shifrlangan tokenlardan foydalaning."),
        ("Whitelist specific trusted origins.", "Faqat ishonchli origin'larni whitelist qiling."),
        ("Do not reflect arbitrary origins. Whitelist trusted domains only.", "Ixtiyoriy origin'larni aks ettirmang, faqat ishonchli domenlarni whitelist qiling."),
        ("Return identical responses for valid and invalid usernames.", "Mavjud va mavjud bo'lmagan foydalanuvchilar uchun bir xil javob qaytaring."),
        ("Use generic error message: 'Invalid username or password'.", "Umumiy xabar qaytaring: 'Foydalanuvchi nomi yoki parol noto'g'ri'."),
        ("Use constant-time comparison for username lookups.", "Username tekshiruvlarida vaqtga bog'liq farq bermaydigan solishtirish usulidan foydalaning."),
        ("Change all default credentials immediately.", "Barcha standart credentiallarni darhol almashtiring."),
        ("Disable debug mode in production. Set DEBUG=False.", "Production muhitida debug rejimini o'chiring va `DEBUG=False` qilib sozlang."),
        ("Disable debug mode. Use custom error pages. Never show stack traces.", "Debug rejimini o'chiring, maxsus error sahifalardan foydalaning va stack trace'larni foydalanuvchiga ko'rsatmang."),
        ("Never expose credentials in API responses. Use environment variables.", "API javoblarida credentiallarni chiqarmang va ularni environment variable'larda saqlang."),
        ("Use bcrypt/argon2 for password hashing. Never expose hashes.", "Parollar uchun bcrypt/argon2 ishlating va hash qiymatlarini oshkor etmang."),
        ("Regenerate session ID after every authentication event.", "Har bir autentifikatsiyadan so'ng session ID'ni yangilang."),
        ("Disable SSLv3/TLS1.0/1.1. Remove weak ciphers. Use TLS 1.2+.", "SSLv3/TLS1.0/TLS1.1 ni o'chiring, zaif cipher'larni olib tashlang va TLS 1.2+ ishlating."),
        ("Enforce authorization at every path level recursively.", "Har bir path darajasida authorization tekshiruvini majburiy qo'llang."),
        ("Enforce authentication on every endpoint. Never trust missing auth.", "Har bir endpointda autentifikatsiyani majburiy tekshiring va auth yo'qligiga ishonib qolmang."),
        ("Restrict HTTP methods. Apply ACL per method.", "HTTP methodlarni cheklang va har bir method uchun ACL ni alohida qo'llang."),
        ("Apply authorization checks on all paths. Use allowlist.", "Barcha path'larda authorization tekshiruvini qo'llang va allowlist asosida ishlang."),
        ("Validate all numeric inputs server-side. Reject negative/zero values.", "Barcha sonli qiymatlarni server tomonda tekshirib, manfiy yoki nol qiymatlarni rad eting."),
        ("Validate auth/authz on every WebSocket message.", "Har bir WebSocket xabari uchun auth/authz tekshiruvini qo'llang."),
    ]
    localized = raw
    for src, dst in replacements:
        localized = localized.replace(src, dst)

    english_markers = [
        "Use ", "Never ", "Disable ", "Validate ", "Whitelist ", "Restrict ",
        "Implement ", "Return ", "Move ", "Fix ", "Apply ", "Do not ",
    ]
    if any(marker in localized for marker in english_markers):
        return _default_recommendation_for_finding(finding)
    return localized


def _recommendation_for_finding(finding: object) -> str:
    raw = _clean_report_text(_attr(finding, "remediation", ""))
    if raw:
        localized = _localize_remediation_text(raw, finding)
        if localized:
            return _shorten(localized, 320)
    return _shorten(_default_recommendation_for_finding(finding), 320)


def _evidence_caption_for_finding(index: int, finding: Optional[object]) -> str:
    figure_no = index + 2
    if not finding:
        return f"{figure_no}-rasm. Ushbu bo'lim uchun dalil rasmi kiritilmagan."
    title = _shorten(_localized_title(finding), 80)
    url = _shorten(_clean_report_text(_attr(finding, "url", "")), 70)
    if url:
        return f"{figure_no}-rasm. \"{title}\" zaifligi bo'yicha dalil skrinshoti ({url})."
    return f"{figure_no}-rasm. \"{title}\" zaifligi bo'yicha dalil skrinshoti."


def _format_finding_paragraph(finding: object, index: int) -> str:
    letter = chr(ord("A") + (index % 26))
    title = _shorten(_localized_title(finding), 180) or "Zaiflik"
    url = _shorten(_clean_report_text(_attr(finding, "url", "")), 180) or "Aniqlanmadi"
    evidence = _evidence_text_for_finding(finding)
    param = _shorten(_clean_report_text(_attr(finding, "param", "")), 80)
    exploit = _shorten(_clean_report_text(_attr(finding, "exploit_cmd", "")), 240)
    recommendation = _recommendation_for_finding(finding)
    parts = [
        f"{letter}) Ekspertiza jarayonida \"{title}\" zaifligi aniqlandi.",
        f"Zaiflik manzili: {url}.",
    ]
    if param:
        parts.append(f"Parametr: {param}.")
    parts.append(f"Zaiflik darajasi: {_risk_uz_label(_attr(finding, 'risk', ''))}.")
    parts.append(f"Dalil: {evidence}")
    if exploit:
        parts.append(f"Ekspluatatsiya/PoC: {exploit}")
    parts.append(f"Tavsiya: {recommendation}")
    return " ".join(parts)


def _collect_remediations(findings: Iterable[object]) -> str:
    seen = []
    for finding in findings:
        remediation = _recommendation_for_finding(finding)
        if remediation and remediation not in seen:
            seen.append(remediation)
    if not seen:
        return "Tavsiyalar: aniqlangan zaifliklarni ustuvorlik bo‘yicha bartaraf etish va himoya sozlamalarini kuchaytirish tavsiya etiladi."
    if len(seen) == 1:
        return f"Tavsiyalar: {seen[0]}"
    return f"Tavsiyalar: {'; '.join(seen[:3])}"


def _virus_status(findings: Iterable[object], report_date: str) -> dict:
    malware_keywords = ["virus", "malware", "webshell", "trojan", "backdoor"]
    infected = []
    for finding in findings:
        haystack = " ".join(
            str(_attr(finding, field, "")).lower()
            for field in ("title", "evidence", "tool", "payload")
        )
        if any(keyword in haystack for keyword in malware_keywords):
            infected.append(finding)
    if infected:
        return {
            "checked": 1,
            "infected": 1,
            "text": (
                f"Veb-saytda {report_date} y. holatiga ko‘ra zararli kod yoki virus alomatlari "
                f"aniqlandi, qo‘shimcha tahlil o‘tkazish tavsiya etiladi"
            ),
        }
    return {
        "checked": 1,
        "infected": 0,
        "text": f"Veb-saytda {report_date} y. holatiga ko‘ra veb-saytda viruslar aniqlanmadi",
    }


def _populate_url_table(table, findings: Iterable[object], fallback_label: str = "Aniqlanmadi") -> None:
    items = list(findings)
    _clear_rows_except_header(table)
    if not items:
        row = table.add_row()
        row.cells[0].text = "1"
        row.cells[1].text = fallback_label
        row.cells[2].text = "-"
        return

    for idx, finding in enumerate(items, start=1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = str(_attr(finding, "url", "") or "Aniqlanmadi")
        row.cells[2].text = _risk_uz_label(str(_attr(finding, "risk", "")))


def _fill_paragraphs(doc: Document, context: dict) -> None:
    target = context["target_parts"]
    start_dt = context["start_dt"]
    end_dt = context["end_dt"]
    end_short = _format_short_date(end_dt)
    intro_text = (
        "“Kiberxavfsizlik markazi” davlat unitar korxonasi \n"
        "va Oʻzbekiston Respublikasi Prezidenti Administratsiyasi huzuridagi taʼlim sifatini "
        f"taʼminlash milliy agentligi oʻrtasida imzolangan ikki tomonlama shartnomaga asosan, "
        f"{_format_long_uzbek_date(start_dt)}dan {_format_long_uzbek_date(end_dt)}ga qadar "
        f"“{target['www_domain']}” (1-rasm) rasmiy veb-sayti kiberxavfsizlik talablariga "
        "muvofiqligi bo‘yicha ekspertizadan o‘tkazildi (OWASP Top-10 \n"
        "va Common Weakness Enumeration zaifliklar ro‘yxati asosida)."
    )
    _set_paragraph_text(_anchor(context, "cover_site"), f"{target['www_domain']} rasmiy veb-sayti")
    _set_paragraph_text(_anchor(context, "intro_body"), intro_text)
    _set_paragraph_text(
        _anchor(context, "intro_caption"),
        f"1-rasm. “{target['www_domain']}” veb-saytining {end_short} y. holatiga ko‘ra asosiy sahifasi ko‘rinishi",
    )
    _set_paragraph_text(_anchor(context, "domain_heading"), f"4. {target['base_domain']} * domen haqida ma’lumot.")
    _set_paragraph_text(
        _anchor(context, "results_heading"),
        f"5. {target['base_domain']} veb-sayti tekshiruvi natijalari bo‘yicha ma’lumot.",
    )


def _fill_tables(doc: Document, context: dict) -> None:
    whois: DomainWhoisInfo = context["whois"]
    target = context["target_parts"]
    tech = context["tech"]
    start_dt = context["start_dt"]
    end_dt = context["end_dt"]
    findings = context["findings"]
    counts = severity_counts(findings)
    report_date = _format_short_date(end_dt)
    virus = context["virus"]

    domain_table = doc.tables[1]
    _update_table_cell(_table_row_by_label(domain_table, "Domen:").cells[1], f"{target['base_domain']} ({target['www_domain']})")
    _update_table_cell(_table_row_by_label(domain_table, "NS server haqida ma’lumot:").cells[1], whois.ns_joined() or "Aniqlanmadi")
    _update_table_cell(_table_row_by_label(domain_table, "Registrator:").cells[1], whois.registrar or "Aniqlanmadi")
    _update_table_cell(_table_row_by_label(domain_table, "Yaratilgan sana:").cells[1], whois.created_date or "Aniqlanmadi")
    _update_table_cell(_table_row_by_label(domain_table, "Yaroqlilik muddati:").cells[1], whois.expiry_date or "Aniqlanmadi")
    _update_table_cell(_table_row_by_label(domain_table, "Veb-server:").cells[1], _friendly_tech_value(tech.get("server"), upper=True))
    _update_table_cell(_table_row_by_label(domain_table, "Dasturlash tili").cells[1], _friendly_tech_value(tech.get("lang")))
    framework = tech.get("framework") or tech.get("cms")
    _update_table_cell(_table_row_by_label(domain_table, "Veb-freymvork/CMS").cells[1], _friendly_tech_value(framework))

    scan_table = doc.tables[2]
    _update_table_cell(_table_row_by_label(scan_table, "Tekshiruv boshlanishi:").cells[1], _format_short_date(start_dt))
    _update_table_cell(_table_row_by_label(scan_table, "Tekshiruv yakunlanishi:").cells[1], report_date)
    _update_table_cell(
        _table_row_by_label(scan_table, "O‘tkazilgan vaqt:").cells[1],
        f"{_business_days(start_dt, end_dt)} ish kuni",
    )

    virus_table = doc.tables[3]
    header_text = f"5.2. VIRUSGA TEKSHRILGANLIK TO‘G‘RISIDA MA’LUMOT {virus['infected']}/{virus['checked']}"
    _update_table_cell(virus_table.rows[0].cells[0], header_text)
    _update_table_cell(virus_table.rows[1].cells[0], virus["text"])

    counts_table = doc.tables[4]
    _update_table_cell(_table_row_by_label(counts_table, "Yuqori darajadagi zaifliklar:").cells[1], str(counts["high"]))
    _update_table_cell(_table_row_by_label(counts_table, "O‘rta darajadagi zaifliklar:").cells[1], str(counts["medium"]))
    _update_table_cell(_table_row_by_label(counts_table, "Quyi darajadagi zaifliklar:").cells[1], str(counts["low"]))

    grouped = context["grouped_findings"]
    _populate_url_table(doc.tables[5], [f for f in grouped.get("6.2", []) if _risk_uz_label(_attr(f, "risk")) == "Yuqori"])
    _populate_url_table(doc.tables[6], [f for f in grouped.get("6.2", []) if _risk_uz_label(_attr(f, "risk")) != "Yuqori"])


def _fill_standard_section_with_context(context: dict, key: str, findings: List[object]) -> None:
    placeholder = _anchor(context, SECTION_PLACEHOLDER_MAP[key])
    section_spec = next(section for section in SECTIONS if section.key == key)
    if not findings:
        _set_paragraph_text(placeholder, section_spec.not_found)
        return

    _set_paragraph_text(placeholder, _format_finding_paragraph(findings[0], 0))
    anchor = placeholder
    for idx, finding in enumerate(findings[1:], start=1):
        anchor = _insert_paragraph_after(anchor, _format_finding_paragraph(finding, idx))


def _fill_section_62(context: dict, findings: List[object]) -> None:
    detail_keys = ["sec62_detail_1", "sec62_detail_2", "sec62_detail_3", "sec62_detail_4"]
    detail_paragraphs = [_anchor(context, key) for key in detail_keys]
    detail_anchor = None

    if not findings:
        _set_paragraph_text(detail_paragraphs[0], "Ushbu turdagi zaiflik aniqlanmadi.")
        detail_anchor = detail_paragraphs[0]
        for para in detail_paragraphs[1:]:
            _delete_paragraph(para)
    else:
        for idx, para in enumerate(detail_paragraphs):
            if idx < len(findings):
                _set_paragraph_text(para, _format_finding_paragraph(findings[idx], idx))
                detail_anchor = para
            else:
                _delete_paragraph(para)
        for idx, finding in enumerate(findings[len(detail_paragraphs):], start=len(detail_paragraphs)):
            detail_anchor = _insert_paragraph_after(detail_anchor, _format_finding_paragraph(finding, idx))

    recommendation = _anchor(context, "sec62_recommendation")
    _set_paragraph_text(recommendation, _collect_remediations(findings))

    for extra_key in ["sec62_password_intro", "sec62_password_1", "sec62_password_2", "sec62_password_3", "sec62_password_4"]:
        _delete_paragraph(_anchor(context, extra_key))


def _fill_sections(doc: Document, context: dict) -> None:
    grouped = context["grouped_findings"]
    for section in SECTIONS:
        findings = list(grouped.get(section.key, []))
        if section.key == "6.2":
            _fill_section_62(context, findings)
        else:
            _fill_standard_section_with_context(context, section.key, findings)


def _make_blank_bytes(size: tuple[int, int], image_format: str) -> bytes:
    image = Image.new("RGB", size, color="white")
    buffer = io.BytesIO()
    image.save(buffer, format=image_format)
    return buffer.getvalue()


def _fit_image_to_canvas(image_bytes: bytes, canvas_size: tuple[int, int], image_format: str) -> bytes:
    canvas = Image.new("RGBA", canvas_size, color=(255, 255, 255, 255))
    source = Image.open(io.BytesIO(image_bytes)).convert("RGBA")
    source.thumbnail(canvas_size, Image.Resampling.LANCZOS)
    offset = (
        max((canvas_size[0] - source.size[0]) // 2, 0),
        max((canvas_size[1] - source.size[1]) // 2, 0),
    )
    canvas.paste(source, offset, source)
    buffer = io.BytesIO()
    output = canvas.convert("RGB") if image_format.upper() in {"JPEG", "JPG"} else canvas
    output.save(buffer, format=image_format)
    return buffer.getvalue()


def _patch_docx_media(docx_path: Path, replacements: Dict[str, bytes]) -> None:
    temp_path = docx_path.with_suffix(".tmp")
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename in replacements:
                data = replacements[item.filename]
            dst.writestr(item, data)
    temp_path.replace(docx_path)


def _build_media_replacements(
    docx_path: Path,
    screenshot_path: Optional[Path],
    fallback_cover_path: Optional[Path],
    logger: Optional[Callable[[str], None]],
) -> Dict[str, bytes]:
    replacements: Dict[str, bytes] = {}
    with zipfile.ZipFile(docx_path, "r") as archive:
        for member in [MEDIA_REPLACEMENTS["cover_page"], MEDIA_REPLACEMENTS["intro_screenshot"], *MEDIA_REPLACEMENTS["blank_slots"]]:
            raw = archive.read(member)
            image = Image.open(io.BytesIO(raw))
            canvas_size = image.size
            image_format = image.format or Path(member).suffix.lstrip(".").upper() or "PNG"

            if member == MEDIA_REPLACEMENTS["cover_page"]:
                if fallback_cover_path and fallback_cover_path.exists():
                    replacements[member] = fallback_cover_path.read_bytes()
                    _log(logger, f"Report: cover page image copied exactly from {fallback_cover_path.name}")
                else:
                    _log(logger, "Report: fallback cover image missing, template cover left unchanged.")
            elif member == MEDIA_REPLACEMENTS["intro_screenshot"]:
                if screenshot_path and screenshot_path.exists():
                    replacements[member] = _fit_image_to_canvas(screenshot_path.read_bytes(), canvas_size, image_format)
                    _log(logger, f"Report: homepage screenshot inserted: {screenshot_path.name}")
                else:
                    _log(logger, "Report: homepage screenshot missing, intro placeholder left unchanged.")
            else:
                replacements[member] = _make_blank_bytes(canvas_size, image_format)
    return replacements


def _convert_with_libreoffice(docx_path: Path, pdf_path: Path) -> bool:
    for binary in ("soffice", "libreoffice"):
        exe = shutil.which(binary)
        if not exe:
            continue
        cmd = [
            exe,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
        if result.returncode == 0:
            generated = pdf_path.parent / f"{docx_path.stem}.pdf"
            if generated.exists():
                if generated != pdf_path:
                    generated.replace(pdf_path)
                return True
    return False


def _convert_with_word(docx_path: Path, pdf_path: Path) -> bool:
    script = f"""
$word = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $doc = $word.Documents.Open('{str(docx_path).replace("'", "''")}')
  $doc.ExportAsFixedFormat('{str(pdf_path).replace("'", "''")}', 17)
  $doc.Close($false)
  $word.Quit()
  exit 0
}} catch {{
  if ($word -ne $null) {{ $word.Quit() }}
  Write-Error $_
  exit 1
}}
"""
    result = subprocess.run(
        ["powershell", "-NoProfile", "-Command", script],
        capture_output=True,
        text=True,
        timeout=240,
    )
    return result.returncode == 0 and pdf_path.exists()


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path, logger: Optional[Callable[[str], None]] = None) -> Optional[Path]:
    if _convert_with_libreoffice(docx_path, pdf_path):
        return pdf_path
    if os.name == "nt" and _convert_with_word(docx_path, pdf_path):
        return pdf_path
    _log(logger, "Report: PDF export skipped because no supported converter was available.")
    return None


def generate_authorized_report(
    *,
    target: str,
    findings: List[object],
    report_dir: str | Path,
    meta: Optional[dict] = None,
    template_path: str | Path | None = None,
    logger: Optional[Callable[[str], None]] = None,
) -> Dict[str, Optional[Path]]:
    meta = meta or {}
    report_root = Path(report_dir)
    report_root.mkdir(parents=True, exist_ok=True)

    template = resolve_template_path(template_path or meta.get("template_path"))
    if not template:
        _log(logger, "Report: DOCX template not found, skipping DOCX/PDF generation.")
        return {"docx": None, "pdf": None}

    target_parts = _extract_target_parts(target)
    safe_target = _safe_name(target_parts["base_domain"])
    timestamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_docx = report_root / f"final_report_{safe_target}_{timestamp}.docx"
    output_pdf = report_root / f"final_report_{safe_target}_{timestamp}.pdf"

    start_dt = _parse_datetime(meta.get("scan_started_at")) or dt.datetime.now()
    end_dt = _parse_datetime(meta.get("scan_finished_at")) or dt.datetime.now()
    tech = dict(meta.get("site_tech") or {})
    grouped_findings = group_findings_by_section(findings)
    whois = lookup_domain_info(target_parts["base_domain"])
    virus = meta.get("virus_status") or _virus_status(findings, _format_short_date(end_dt))
    fallback_cover_path = resolve_cover_image_path(meta.get("default_cover_path"))
    screenshot_source_url = _resolve_screenshot_url(target_parts, logger)

    with tempfile.TemporaryDirectory(prefix="deepsec_report_") as temp_dir:
        screenshot_path = capture_homepage_screenshot(
            screenshot_source_url,
            Path(temp_dir) / "homepage.png",
            logger=logger,
        )

        shutil.copy2(template, output_docx)
        doc = Document(str(output_docx))
        anchors = _build_anchor_map(doc)
        context = {
            "anchors": anchors,
            "target_parts": target_parts,
            "start_dt": start_dt,
            "end_dt": end_dt,
            "tech": tech,
            "whois": whois,
            "virus": virus,
            "findings": findings,
            "grouped_findings": grouped_findings,
        }
        _fill_paragraphs(doc, context)
        _fill_tables(doc, context)
        _fill_sections(doc, context)

        sec62_findings = list(grouped_findings.get("6.2", []))
        _set_paragraph_text(_anchor(context, "sec62_caption_2"), _evidence_caption_for_finding(0, sec62_findings[0] if len(sec62_findings) > 0 else None))
        _set_paragraph_text(_anchor(context, "sec62_caption_3"), _evidence_caption_for_finding(1, sec62_findings[1] if len(sec62_findings) > 1 else None))
        _set_paragraph_text(_anchor(context, "sec62_caption_4"), _evidence_caption_for_finding(2, sec62_findings[2] if len(sec62_findings) > 2 else None))
        _set_paragraph_text(_anchor(context, "sec62_caption_5"), _evidence_caption_for_finding(3, sec62_findings[3] if len(sec62_findings) > 3 else None))
        doc.save(str(output_docx))

        replacements = _build_media_replacements(output_docx, screenshot_path, fallback_cover_path, logger)
        if replacements:
            _patch_docx_media(output_docx, replacements)

    pdf_result = convert_docx_to_pdf(output_docx, output_pdf, logger=logger)
    return {"docx": output_docx, "pdf": pdf_result}
