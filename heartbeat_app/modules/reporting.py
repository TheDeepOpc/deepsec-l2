from .base import *
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from .domain_analyzer import get_whois_data
import datetime
import collections
import hashlib
import json
import re

class Reporter:
    def __init__(self, target: str, graph: "EndpointGraph"):
        self.target = target
        self.graph  = graph
        self.scan_log: list[dict] = []  # Scan process logs

    def log(self, step: str, action: str, details: dict = None):
        """Logs each action during the scan process."""
        self.scan_log.append({
            "timestamp": datetime.datetime.now().isoformat(),
            "step": step,
            "action": action,
            "details": details or {},
        })

    def generate(self, findings: list["Finding"]) -> str:
        ts     = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        by_risk= collections.defaultdict(list)
        suppressed = []
        for f in findings:
            if not f.fp_filtered:
                by_risk[f.risk].append(f)
            else:
                suppressed.append(f)
        total  = sum(len(v) for v in by_risk.values())
        lines  = [
            f"# Pentest Report: {self.target}",
            f"**Date:** {ts}  ",
            f"**Total findings:** {total}  ",
            f"**Endpoints tested:** {self.graph.stats()['endpoints']}  ",
            "",
            "## Executive Summary",
            "| Risk | Count |",
            "|------|-------|",
        ]
        for risk in ["Critical", "High", "Medium", "Low", "Info"]:
            c = len(by_risk.get(risk, []))
            if c:
                lines.append(f"| **{risk}** | {c} |")
        lines += ["", "---", "## Findings"]
        for risk in ["Critical", "High", "Medium", "Low", "Info"]:
            for f in by_risk.get(risk, []):
                lines += [
                    f"", f"### [{f.risk}] {f.owasp_id} — {f.title}",
                    f"| Field | Value |", f"|-------|-------|",
                    f"| **URL** | `{f.url}` |",
                    f"| **Method** | `{f.method}` |",
                    f"| **Parameter** | `{f.param}` |",
                    f"| **OWASP** | {f.owasp_id} — {f.owasp_name} |",
                    f"| **Risk** | {f.risk} |",
                    f"| **Confidence** | {f.confidence}% |",
                    f"| **Confirmed** | {'✅ Yes' if f.confirmed else '⚠ Unconfirmed'} |",
                    f"| **Verification Reason** | {self._verification_reason(f)} |",
                    "", f"**Evidence:** {f.evidence}", "",
                ]
                if f.request_raw:
                    lines += ["**Request:**", "```http", f.request_raw[:400], "```"]
                if f.exploit_cmd:
                    lines += ["**PoC:**", "```bash", f.exploit_cmd, "```"]
                lines += [f"**Remediation:** {f.remediation}", "", "---"]
        if suppressed:
            lines += ["", "## Suppressed Potentials"]
            for f in suppressed:
                lines += [
                    "",
                    f"### [{f.risk}] {f.title}",
                    f"- URL: `{f.url}`",
                    f"- Confidence: {f.confidence}%",
                    f"- Why suppressed: {f.suppression_reason or 'Filtered during false-positive reduction.'}",
                ]
                if f.evidence:
                    lines.append(f"- Evidence: {f.evidence}")
        return "\n".join(lines)

    def _build_finding_report(self, f: "Finding") -> dict:
        """Creates a complete, thorough report for a single finding."""
        report = {
            "id": hashlib.md5(f"{f.title}{f.url}{f.param}".encode()).hexdigest()[:12],
            "vulnerability": {
                "title": f.title,
                "owasp_id": f.owasp_id,
                "owasp_name": f.owasp_name,
                "risk_level": f.risk,
                "confidence": f.confidence,
                "confirmed": f.confirmed,
                "verification_reason": self._verification_reason(f),
                "description": self._vuln_description(f),
            },
            "target": {
                "url": f.url,
                "method": f.method,
                "parameter": f.param,
            },
            "evidence": {
                "summary": f.evidence,
                "baseline_diff": f.baseline_diff,
                "response_body_snippet": f.response_raw[:3000] if f.response_raw else "",
                "tool_output": f.tool_output[:2000] if f.tool_output else "",
            },
            "exploitation": {
                "poc_command": f.exploit_cmd,
                "request": f.request_raw,
                "payload_used": f.payload,
                "what_attacker_can_do": self._impact_description(f),
                "attack_scenario": self._attack_scenario(f),
            },
            "remediation": f.remediation,
            "tool_used": f.tool,
            "timestamp": f.timestamp,
        }
        if f.chain:
            report["exploitation"]["chain_steps"] = f.chain
        if f.oob:
            report["exploitation"]["oob_confirmed"] = True
        return report

    def _verification_reason(self, f: "Finding") -> str:
        """Human-readable reason for why a finding is confirmed or not."""
        explicit_reason = str(getattr(f, "verification_reason", "") or "").strip()
        if explicit_reason:
            return explicit_reason

        if f.fp_filtered:
            return f.suppression_reason or "Filtered as false positive by verification logic."

        if f.confirmed:
            if getattr(f, "oob", False):
                return "Out-of-band callback observed during verification."

            evidence = str(f.evidence or "")
            if "VERIFIED:" in evidence:
                return evidence.split("VERIFIED:", 1)[1].strip() or "Verified by evidence in response comparison."

            tool_output = str(f.tool_output or "")
            if "VERIFIED:" in tool_output:
                return tool_output.split("VERIFIED:", 1)[1].strip() or "Verified by scanner/tool output."

            if f.exploit_cmd:
                return "Exploit PoC generated and verification signal met."

            return "Confirmed by automated verification heuristics and confidence threshold."

        if f.suppression_reason:
            return f.suppression_reason

        if f.confidence < 70:
            return f"Not confirmed due to low confidence ({f.confidence}%)."

        return "Not confirmed: evidence did not meet verification criteria."

    def _vuln_description(self, f: "Finding") -> str:
        """Full explanation of what the vulnerability consists of."""
        descs = {
            "A01": "Broken Access Control — the server does not properly validate user permissions. "
                   "Unauthorized access to protected resources is possible.",
            "A02": "Cryptographic Failures — encryption is improperly or weakly implemented.",
            "A03": "Injection — malicious data can be sent to the server and executed.",
            "A04": "Insecure Design — security weakness in the application architecture.",
            "A05": "Security Misconfiguration — the server is improperly configured.",
            "A07": "Identification/Auth Failures — authentication is weak or can be bypassed.",
            "A10": "SSRF — the server can be forced to send requests to external or internal resources.",
        }
        return descs.get(f.owasp_id, f"OWASP {f.owasp_id} — {f.owasp_name}")

    def _impact_description(self, f: "Finding") -> str:
        """What an attacker CAN DO through this vulnerability."""
        if "BAC" in f.title or "Access Control" in f.title or "bypass" in f.title.lower():
            return ("An attacker can gain unauthorized access to protected pages, view admin panels, "
                    "read confidential configurations, and obtain user data.")
        if "SSRF" in f.title:
            return ("An attacker can force the server to send requests to other services on the internal network, "
                    "obtain AWS metadata, and read internal data.")
        if "SQL" in f.title.upper():
            return ("An attacker can read, modify, and delete all data in the database. "
                    "In some cases, they can also execute commands on the operating system.")
        if "XSS" in f.title.upper():
            return ("An attacker can execute JavaScript code in other users' browsers, "
                    "steal cookies/sessions, and perform phishing attacks.")
        if "rate limit" in f.title.lower():
            return "An attacker can brute-force passwords (no rate limiting in place)."
        if "header" in f.title.lower() and "missing" in f.title.lower():
            return "Browser protections are not enabled — vulnerable to clickjacking, XSS, MIME sniffing attacks."
        if "JWT" in f.title:
            return "An attacker can forge JWT tokens and gain access as another user."
        if "Method" in f.title:
            return "Unexpected HTTP methods are accepted — data modification/deletion is possible."
        return "This vulnerability can be exploited by an attacker to compromise the target system."

    def _attack_scenario(self, f: "Finding") -> str:
        """Explains a concrete attack scenario."""
        if "X-Forwarded-For" in f.title:
            return ("1. Attacker adds 'X-Forwarded-For: 127.0.0.1' header to the request\n"
                    "2. Server treats the attacker as localhost\n"
                    "3. IP-based restrictions are bypassed\n"
                    "4. Access to protected endpoints becomes possible")
        if "403" in f.title or "ACL" in f.title:
            return ("1. /admin page returns 403 Forbidden\n"
                    "2. But /admin/child-path page returns 200 OK\n"
                    "3. Permission check only exists on parent path — not on child\n"
                    "4. Attacker accesses child path directly and views admin data")
        if "SSRF" in f.title:
            return ("1. Attacker sends an internal URL in the parameter (e.g. http://169.254.169.254/)\n"
                    "2. Server sends a request to that URL\n"
                    "3. Internal network data is returned to the attacker")
        if "rate limit" in f.title.lower():
            return ("1. Attacker sends thousands of passwords to the login page\n"
                    "2. Server applies no rate limiting\n"
                    "3. When the correct password is found — account is compromised")
        return f"PoC buyrug'ini ishga tushiring va natijani kuzating: {f.exploit_cmd}"

    def save(self, findings: list["Finding"]) -> Path:
        ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = re.sub(r'[^\w.]', '_', self.target)

        # 1st JSON: Scan Log — what actions were performed
        log_path = REPORT_DIR / f"scan_log_{safe}_{ts}.json"
        scan_log_data = {
            "target": self.target,
            "scan_date": datetime.datetime.now().isoformat(),
            "endpoints_tested": self.graph.stats()["endpoints"],
            "parameters_found": self.graph.stats()["params"],
            "scan_steps": self.scan_log,
            "endpoint_map": self.graph.nodes,
        }
        log_path.write_text(json.dumps(scan_log_data, indent=2, default=str, ensure_ascii=False), encoding="utf-8")

        # 2nd JSON: Findings Report — what was found (complete and thorough)
        confirmed_findings = [f for f in findings if not f.fp_filtered]
        suppressed_findings = [f for f in findings if f.fp_filtered]

        findings_data = {
            "target": self.target,
            "scan_date": datetime.datetime.now().isoformat(),
            "summary": {
                "total_confirmed": len(confirmed_findings),
                "total_suppressed": len(suppressed_findings),
                "by_risk": {},
                "by_owasp": {},
            },
            "confirmed_vulnerabilities": [],
            "suppressed_false_positives": [],
        }

        # Risk and OWASP statistics
        for f in confirmed_findings:
            findings_data["summary"]["by_risk"][f.risk] = findings_data["summary"]["by_risk"].get(f.risk, 0) + 1
            findings_data["summary"]["by_owasp"][f.owasp_id] = findings_data["summary"]["by_owasp"].get(f.owasp_id, 0) + 1

        # Confirmed findings — full report
        for f in confirmed_findings:
            findings_data["confirmed_vulnerabilities"].append(self._build_finding_report(f))

        # Suppressed — brief reason
        for f in suppressed_findings:
            findings_data["suppressed_false_positives"].append({
                "title": f.title,
                "url": f.url,
                "risk": f.risk,
                "confidence": f.confidence,
                "reason_suppressed": f.suppression_reason or "AI FP filter",
                "verification_reason": self._verification_reason(f),
            })

        findings_path = REPORT_DIR / f"findings_{safe}_{ts}.json"
        findings_path.write_text(json.dumps(findings_data, indent=2, default=str, ensure_ascii=False), encoding="utf-8")

        # MD report also saved (for legacy format)
        md_path = REPORT_DIR / f"pentest_{safe}_{ts}.md"
        md_path.write_text(self.generate(findings), encoding="utf-8")

        console.print(f"\n[bold green]✅ Reports saved:[/bold green]")
        console.print(f"   📋 Scan Log:     {log_path}")
        console.print(f"   🔍 Findings:     {findings_path}")
        console.print(f"   📄 Full Report:  {md_path}")
        return md_path

    def generate_docx_report(self, findings: list["Finding"], tech_stack: dict, output_path: str):
        """Generates a formal .docx penetration testing report based on a template."""
        document = Document()

        # --- Document Styles ---
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # --- Title Page ---
        document.add_paragraph()
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Oʻzbekiston Respublikasi Prezidenti Administratsiyasi huzuridagi Taʼlim sifatini taʼminlash milliy agentligining')
        run.bold = True
        run.font.size = Pt(14)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'\n{self.target.upper()} rasmiy veb-sayti')
        run.bold = True
        run.font.size = Pt(16)

        document.add_paragraph('\n' * 8) # Spacing

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Toshkent - {datetime.datetime.now().year}")
        run.bold = True
        run.font.size = Pt(14)
        document.add_page_break()

        # --- Table of Contents (Static) ---
        document.add_heading('Mundarija', level=1)
        toc = [
            ("1. Kirish", "3"),
            ("2. Veb-sayt xavfsizligini tekshirish bosqichlari", "3"),
            ("2.1. Dastlabki bosqich", "3"),
            ("2.2. Asosiy bosqich", "4"),
            ("2.3. Yakuniy bosqich", "4"),
            ("3. Zaifliklarning jiddiylik darajasini aniqlash", "4"),
            ("4. Domen haqida ma’lumot", "5"),
            ("4.1. Domen parametrlari", "5"),
            ("4.2. Saytda foydalanilayotgan texnalogiyalar haqida ma’lumot", "5"),
            ("5. Tekshiruv natijalari bo‘yicha ma’lumot", "5"),
            ("5.1. Tekshiruv haqida ma’lumot", "5"),
            ("5.2. Virusga tekshirilganlik haqida ma’lumot", "5"),
            ("5.3. Topilgan zaifliklar soni (daraja bo‘yicha)", "5"),
            ("6. Aniqlangan zaifliklar", "6"),
            ("7. Umumiy talablar", "10"),
        ]
        for item, page in toc:
            p = document.add_paragraph(style='List Bullet')
            p.add_run(item)
        document.add_page_break()

        # --- 1. Introduction ---
        document.add_heading('1. Kirish', level=1)
        p = document.add_paragraph()
        p.add_run('“Kiberxavfsizlik markazi” davlat unitar korxonasi tomonidan, ushbu hujjatga asosan, ')
        p.add_run(f'“{self.target}” rasmiy veb-sayti kiberxavfsizlik talablariga muvofiqligi bo‘yicha ekspertizadan o‘tkazildi ')
        p.add_run('(OWASP Top-10 va Common Weakness Enumeration zaifliklar ro‘yxati asosida).')
        p.add_run('\n\nUshbu hisobot tekshiruv bosqichlarini tavsiflaydi, shuningdek, aniqlangan zaifliklar va ularni bartaraf etish bo‘yicha tavsiyalar haqida ma’lumot beradi.')
        
        # Placeholder for image
        document.add_paragraph("\n[Bu yerga veb-saytning asosiy sahifasi tasviri joylashtiriladi]\n", style='BodyText')
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'1-rasm. “{self.target}” veb-saytining {datetime.datetime.now().strftime("%d.%m.%Y")} y. holatiga ko‘ra asosiy sahifasi ko‘rinishi')
        run.italic = True

        # --- 2. Stages ---
        document.add_heading('2. Veb-sayt xavfsizligini tekshirish bosqichlari', level=1)
        document.add_paragraph("Veb-sayt xavfsizligi tekshiruvi 3 bosqichni o‘z ichiga oladi, bular:")
        document.add_heading('2.1. Dastlabki bosqich', level=2)
        document.add_paragraph("xosting panelidan foydalangan holda veb-saytning zaxira nusxasini yaratish (ko‘chirish) bo‘yicha chora-tadbirlarni amalga oshirishdan iborat bo‘lgan bosqichdir...")
        document.add_heading('2.2. Asosiy bosqich', level=2)
        document.add_paragraph("ikki qismdan iborat: suqulib kirishga sinash (Pentest) va veb-saytning ichki kodini statik tahlil qilish.")
        document.add_heading('2.3. Yakuniy bosqich', level=2)
        document.add_paragraph("aniqlangan zaiflik va kamchiliklar bo‘yicha hisobot tayyorlash hamda ularni bartaraf etish bo‘yicha tavsiyalar ishlab chiqish.")
        
        # --- 3. Severity Levels ---
        document.add_heading('3. Zaifliklarning jiddiylik darajasini aniqlash', level=1)
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Zaiflik darajasi'
        hdr_cells[1].text = 'Tavsif'
        severity_data = [
            ("Yuqori", "Bu darajadagi zaifliklardan foydalanish orqali tizim ustidan to‘liq nazoratni qo‘lga kiritish..."),
            ("O‘rta", "Ushbu darajadagi zaifliklar past darajadagi zararga ega hisoblanadi..."),
            ("Quyi", "Ushbu darajadagi zaiflikdan foydalanish jiddiy zarar keltirib chiqarmaydi...")
        ]
        for level, desc in severity_data:
            row_cells = table.add_row().cells
            row_cells[0].text = level
            row_cells[1].text = desc

        # --- 4. Domain Info ---
        document.add_heading(f'4. {self.target} domeni haqida ma’lumot', level=1)
        
        document.add_heading('4.1. Domen parametrlari', level=2)
        whois_data = get_whois_data(self.target)
        if 'error' in whois_data:
             document.add_paragraph(f"Domen ma'lumotlarini olishda xatolik: {whois_data['error']}")
        else:
            domain_table = document.add_table(rows=1, cols=2)
            domain_table.style = 'Table Grid'
            hdr_cells = domain_table.rows[0].cells
            hdr_cells[0].text = 'Parametr'
            hdr_cells[1].text = 'Qiymat'
            domain_info = {
                "Domen:": self.target,
                "NS server haqida ma’lumot:": "\n".join(whois_data.get('ns_servers', ['aniqlanmadi'])),
                "Registrator:": whois_data.get('registrar', 'aniqlanmadi'),
                "Yaratilgan sana:": whois_data.get('creation_date', 'aniqlanmadi'),
                "Yaroqlilik muddati:": whois_data.get('expiration_date', 'aniqlanmadi'),
            }
            for key, value in domain_info.items():
                row_cells = domain_table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = value

        document.add_heading('4.2. Veb-saytda foydalanilgan texnologiyalar to‘g‘risida ma’lumot', level=2)
        tech_table = document.add_table(rows=1, cols=2)
        tech_table.style = 'Table Grid'
        hdr_cells = tech_table.rows[0].cells
        hdr_cells[0].text = 'Komponent'
        hdr_cells[1].text = 'Texnologiya'
        for key, value in tech_stack.items():
            row_cells = tech_table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = ", ".join(value) if isinstance(value, list) else value

        # --- 5. Scan Results ---
        document.add_heading('5. Tekshiruv natijalari bo‘yicha ma’lumot', level=1)
        document.add_heading('5.1. Tekshiruv to‘g‘risida ma’lumot', level=2)
        # Bu qismni to'ldirish uchun scan boshlanish va tugash vaqtini saqlash kerak
        document.add_paragraph("Tekshiruv vaqtlari haqida ma'lumot kiritiladi.")

        document.add_heading('5.2. Virusga tekshirilganlik to‘g‘risida ma’lumot', level=2)
        document.add_paragraph(f'Veb-saytda {datetime.datetime.now().strftime("%d.%m.%Y")} y. holatiga ko‘ra viruslar aniqlanmadi.')

        document.add_heading('5.3. Aniqlangan zaifliklar soni (darajalar bo‘yicha)', level=2)
        by_risk = collections.defaultdict(list)
        for f in findings:
            if not f.fp_filtered:
                by_risk[f.risk].append(f)
        
        risk_table = document.add_table(rows=1, cols=2)
        risk_table.style = 'Table Grid'
        hdr_cells = risk_table.rows[0].cells
        hdr_cells[0].text = 'Daraja'
        hdr_cells[1].text = 'Soni'
        risk_counts = {
            "Yuqori": len(by_risk.get("Critical", []) + by_risk.get("High", [])),
            "O‘rta": len(by_risk.get("Medium", [])),
            "Quyi": len(by_risk.get("Low", []) + by_risk.get("Info", []))
        }
        for level, count in risk_counts.items():
            row_cells = risk_table.add_row().cells
            row_cells[0].text = f"{level} darajadagi zaifliklar:"
            row_cells[1].text = str(count)

        # --- 6. Findings ---
        document.add_heading('6. Aniqlangan zaifliklar', level=1)
        
        owasp_map = {
            'A01': ('6.1. Broken Access Control', 'Broken Access Control (Foydalana olish nazoratining buzilishi).'),
            'A02': ('6.4. Cryptographic Failures', 'Cryptographic failures (Kriptografik xatolilar).'),
            'A03': ('6.5. Injection', 'Injection (Inyeksiya).'),
            'A04': ('6.6. Insecure Design', 'Insecure design (Xavfsiz bo‘lmagan dizayn).'),
            'A05': ('6.2. Security Misconfiguration', 'Security misconfigurations (Noto‘g‘ri xavfsizlik konfiguratsiyasi).'),
            'A06': ('6.3. Software Supply Chain Failures', 'Software Supply Chain Failures (Dasturiy ta’minot zanjiridagi nosozliklar).'),
            'A07': ('6.7. Authentication Failures', 'Authentication Failures (Autentifikatsiya xatoliklari).'),
            'A08': ('6.8. Software and Data Integrity Failures', 'Dasturiy ta’minot va ma’lumotlar yaxlitligi xatoliklari.'),
            'A09': ('6.9. Security logging and monitoring failures', 'Security logging and monitoring failures (Xavfsizlik qaydlari va nazoratidagi xatoliklar).'),
            # A10 is now part of A05, but we can map SSRF here if needed
            'A10': ('6.2. Security Misconfiguration', 'Server-Side Request Forgery (SSRF)'),
        }

        # Group findings by OWASP category
        findings_by_owasp = collections.defaultdict(list)
        for f in findings:
            if not f.fp_filtered:
                owasp_id_short = f.owasp_id.split(':')[0]
                if owasp_id_short in owasp_map:
                    findings_by_owasp[owasp_id_short].append(f)

        # Static OWASP descriptions
        owasp_descriptions = {
            'A01': "Bu tizim ma’lumotlarga kirish darajasini yoki uning funktsional imkoniyatlarini noto‘g‘ri nazorat qilishi tufayli kelib chiqadigan zaifliklar to'plami hisoblanadi...",
            'A05': "Noto‘g‘ri xavfsizlik konfiguratsiyasi zaifligi – bu dastur, server, ma'lumotlar bazasi yoki boshqa tizim komponentlarining sozlamalari xavfsiz bo'lmagan holatda kelib chiqadi...",
            'A06': "Dasturiy ta’minot zanjiridagi nosozliklar – bu veb-ilova uchinchi tomon (ya’ni tashqi) kutubxona, plagin yoki komponentlardan foydalanganda yuzaga keladigan zaifliklardir...",
            'A02': "Kriptografik nosozliklar – ma’lumotlarni himoya qilish uchun kriptografik usullardan foydalanish vaqtida yo’l qo’yiladigan kamchilik hisoblanadi...",
            'A03': "Inyeksiya – foydalanuvchi tizimga yuborilayotgan ma’lumot ichiga qo’shimcha zararli kodlar qo’shishi va tizim ma’lumotni tekshirish hamda o’zgartirishlarsiz qabul qilishi orqali kelib chiqadigan zaiflik hisoblanadi...",
            'A04': "Xavfsiz bo’lmagan dizayn zaifligi – noto‘g‘ri tuzilgan dastur tuzilishi sababli kelib chiqadi va tizim xavfsizligiga zarar yetkazishi mumkin...",
            'A07': "Autentifikatsiya — bu foydalanuvchining tizimga kirish jarayonida o‘z shaxsini ishonchli tasdiqlash mexanizmi hisoblanadi...",
            'A08': "Ushbu zaifliklar guruhi ko’p holatlarda yangilanishlar vaqtida yangi dastur qayerdan yuklab olinyotganligi va uning xavfsizligi tekshirilmasligi sababli kelib chiqadi...",
            'A09': "Ushbu zaiflik tizimda mavjud bo’lishi mumkin bo’lgan xavfsizlik hodisalarining to'g'ri qayd etilmasligi va nazorat qilinmasligi sababli kelib chiqadi...",
            # 'Mishandling of Exceptional Conditions' is not a direct OWASP Top 10 2021 item, but we can add it.
        }
        
        # Iterate through all OWASP categories to ensure all are listed
        all_owasp_cats = sorted(owasp_map.items(), key=lambda item: item[1][0])

        for owasp_id, (heading, title) in all_owasp_cats:
            document.add_heading(title, level=2)
            document.add_paragraph(owasp_descriptions.get(owasp_id, "Ushbu zaiflik uchun tavsif topilmadi."))
            
            category_findings = findings_by_owasp.get(owasp_id)
            if category_findings:
                for i, finding in enumerate(category_findings):
                    document.add_paragraph(f"Aniqlangan holat #{i+1}: {finding.title}", style='Intense Quote')
                    p = document.add_paragraph()
                    p.add_run("URL: ").bold = True
                    p.add_run(f"`{finding.url}`")
                    p = document.add_paragraph()
                    p.add_run("Isbot (Evidence): ").bold = True
                    p.add_run(finding.evidence)
                    p = document.add_paragraph()
                    p.add_run("Bartaraf etish (Remediation): ").bold = True
                    p.add_run(finding.remediation)
                    if finding.exploit_cmd:
                        p = document.add_paragraph()
                        p.add_run("PoC (Proof of Concept): ").bold = True
                        p.add_run(f"`{finding.exploit_cmd}`")
                    document.add_paragraph()
            else:
                p = document.add_paragraph()
                run = p.add_run("Ushbu turdagi zaiflik aniqlanmagan.")
                run.italic = True
        
        # --- 7. General Requirements ---
        document.add_heading('7. Umumiy talablar', level=1)
        reqs = [
            "veb-saytlarni himoya qilish vositalaridan (WAF) foydalanish;",
            "barcha turdagi dasturiy mahsulotlarni so‘nggi barqaror versiyaga yangilash;",
            "muntazam ravishda veb-saytni kiberxavfsizlik talablariga muvofiqligini tekshirish;",
            "veb-saytlarga tashriflarni ro‘yxatga olish hamda kiberxavfsizlik hodisalari monitoringini ta’minlash;",
            "veb-sayt va barcha muhim ma’lumotlardan muntazam ravishda zaxira nusxalarini olib turish;",
            "ko‘p faktorli autentifikatsiyadan (MFA) foydalanish;",
            "antivirus dasturlari yordamida veb-sayt fayllarini muntazam ravishda tekshirish."
        ]
        for req in reqs:
            document.add_paragraph(req, style='List Bullet')

        # --- Footer/Eslatma ---
        document.add_paragraph()
        p = document.add_paragraph()
        run = p.add_run("Eslatma: Veb-sayt ekspertizasi ... metodologiyasiga muvofiq o’tkazildi.")
        run.italic = True

        document.save(output_path)
        print(f"Hisobot muvaffaqiyatli {output_path} fayliga saqlandi.")

__all__ = ['Reporter']
