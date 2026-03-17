from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell
import datetime
import os
from .domain_analyzer import get_whois_data

# Assuming domain_analyzer.py is in the same directory or accessible
# from heartbeat_app.modules.domain_analyzer import DomainAnalyzer

def set_cell_color(cell: _Cell, color: str):
    """
    Set cell background color.
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

class ReportGenerator:
    def __init__(self, target_domain, report_data):
        self.target_domain = target_domain
        self.report_data = report_data
        self.document = Document()
        self.setup_styles()

    def setup_styles(self):
        """Sets up the default styles for the document."""
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

    def add_title_page(self):
        """Adds the title page to the document."""
        # Clean up previous paragraphs if any
        for para in self.document.paragraphs:
            p = para._element
            p.getparent().remove(p)

        # Add logo
        logo_path = os.path.join('report_settings', 'logo.png')
        if os.path.exists(logo_path):
            self.document.add_picture(logo_path, width=Inches(2.0))
        else:
            self.document.add_paragraph("Logo not found")

        # Title
        title = self.document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("O‘zbekiston Respublikasi Prezidenti\nAdministratsiyasi huzuridagi Ta’lim sifatini\nta’minlash milliy agentligining\n")
        run.bold = True
        run.font.size = Pt(14)
        
        run = title.add_run(f"“{self.target_domain}” rasmiy veb-sayti\n")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

        title.add_run("KIBERXAVFSIZLIK TALABLARIGA\nMUVOFIQLIGI YUZASIDAN\nEKSPERTIZA HISOBOTI")
        
        # Add space
        self.document.add_paragraph()
        self.document.add_paragraph()

        # Footer-like info at the bottom of the page
        # This is a simple implementation. For a real footer, docx section footers should be used.
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run("+998 71 203 55 11\t\t")
        p.add_run("T. Shevchenko, 20\t\t")
        p.add_run("info@csec.uz\t\t")
        p.add_run("www.csec.uz")

    def add_table_of_contents(self):
        self.document.add_page_break()
        self.document.add_heading('MUNDARIJA', level=1)
        toc = [
            ("1. Kirish", 3),
            ("2. Veb-sayt xavfsizligini tekshirish bosqichlari", 3),
            ("2.1. Dastlabki bosqich", 3),
            ("2.2. Asosiy bosqich", 4),
            ("2.3. Yakuniy bosqich", 4),
            ("3. Zaifliklarning jiddiylik darajasini aniqlash", 4),
            ("4. Domen haqida ma’lumot", 5),
            ("4.1. Domen parametrlari", 5),
            ("4.2. Saytda foydalanilayotgan texnalogiyalar haqida ma’lumot", 5),
            ("5. Tekshiruv natijalari bo‘yicha ma’lumot", 5),
            ("5.1. Tekshiruv haqida ma’lumot", 5),
            ("5.2. Virusga tekshirilganlik haqida ma’lumot", 5),
            ("5.3. Topilgan zaifliklar soni (daraja bo‘yicha)", 5),
            ("6. Aniqlangan zaifliklar", 6),
            ("6.1. Broken Access Control", 6),
            ("6.2. Security Misconfiguration", 6),
            ("6.3. Software Supply Chain Failures", 6),
            ("6.4. Cryptographic Failures", 7),
            ("6.5. Injection (Inyeksiya)", 9),
            ("6.6. Insecure Design", 9),
            ("6.7. Authentication Failures", 9),
            ("6.8. Software and Data Integrity Failures", 10),
            ("6.9. Security logging and monitoring failures", 10),
            ("6.10. Mishandling of Exceptional Conditions", 10),
            ("7. Umumiy talablar", 10),
        ]
        for item, page in toc:
            p = self.document.add_paragraph()
            p.add_run(item)
            p.add_run().add_tab()
            p.add_run(str(page))
            
    def add_introduction(self):
        self.document.add_page_break()
        self.document.add_heading('1. Kirish', level=1)
        p = self.document.add_paragraph()
        p.add_run('“Kiberxavfsizlik markazi” davlat unitar korxonasi va Oʻzbekiston Respublikasi Prezidenti Administratsiyasi huzuridagi taʼlim sifatini taʼminlash milliy agentligi oʻrtasida imzolangan ikki tomonlama shartnomaga asosan, “')
        p.add_run(f'{self.target_domain}').bold = True
        p.add_run('” (1-rasm) rasmiy veb-sayti kiberxavfsizlik talablariga muvofiqligi bo‘yicha ekspertizadan o‘tkazildi (OWASP Top-10 va Common Weakness Enumeration zaifliklar ro‘yxati asosida).')
        self.document.add_paragraph('Ushbu hisobot tekshiruv bosqichlarini tavsiflaydi, shuningdek, aniqlangan zaifliklar va ularni bartaraf etish bo‘yicha tavsiyalar haqida ma’lumot beradi.')
        self.document.add_paragraph("Bu yerda rasm bo'ladi") # Placeholder for image
        self.document.add_paragraph('Veb-sayt xavfsizligi tekshiruvi mavjud xavflarni oldini olish, hamda kiberxavfsizlik zaifliklarini aniqlash, shuningdek ularni bartaraf etish maqsadida amalga oshirildi. Bundan tashqari, veb-saytni tekshirishda quyidagilar amalga oshirildi:')
        self.document.add_paragraph('- veb-saytdagi taniqli va noma’lum “Zero day” zaifliklarini tahlil qilish va aniqlash;')
        self.document.add_paragraph('- mavjud kiberxavfsizlik tahdidlarini aniqlash;')
        self.document.add_paragraph('- aniqlangan zaifliklarni samarali bartaraf etish choralarini ko‘rish bo‘yicha tavsiyalar ishlab chiqish.')

    def add_vulnerability_severity_levels(self):
        self.document.add_heading('3. Zaifliklarning jiddiylik darajasini aniqlash', level=1)
        p = self.document.add_paragraph('Quyidagi jadval ushbu hujjatda kiberxavfsizlik tahdidlarini baholash uchun foydalanilgan zaiflik darajalarini tavsiflaydi.')
        
        table = self.document.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        levels = [
            ("Yuqori darajadagi zaifliklar:", "FF0000", "Bu darajadagi zaifliklardan foydalanish orqali tizim ustidan to‘liq nazoratni qo‘lga kiritish, muhim huquqlarga ega bo‘lish, soxta ma’lumotlar tarqatilishi hamda jiddiy ma’lumotlar o‘g‘irlanishi mumkin bo‘ladi.\nTavsiya: qisqa vaqt ichida rejalashtirish, hamda zaifliklarni bartaraf etish tavsiya qilinadi."),
            ("O‘rta darajadagi zaifliklar:", "FFFF00", "Ushbu darajadagi zaifliklar past darajadagi zararga ega hisoblanadi yoki zaiflikdan foydalanish uchun qo‘shimcha ma’lumot va resurslar talab qilinadi.\nTavsiya: yuqori darajadagi zaifliklar bartaraf etilgan so‘ng, rejalashtirish hamda zaifliklarni bartaraf etish tavsiya qilinadi."),
            ("Quyi darajadagi zaifliklar:", "00FF00", "Ushbu darajadagi zaiflikdan foydalanish jiddiy zarar keltirib chiqarmaydi yoki foydalanish uchun qo‘shimcha ma’lumot va resurslar talab qilinadi. Zaiflikdan foydalanish tizim haqida qo‘shimcha ma’lumotlarni yig‘ishga yordam berishi mumkin.\nTavsiya: Oxirgi o‘rinlarda rejalashitirish hamda zaifliklarni bartaraf etish tavsiya qilinadi.")
        ]

        for i, (level, color, desc) in enumerate(levels):
            cell1 = table.cell(i, 0)
            cell1.text = level
            set_cell_color(cell1, color)
            
            cell2 = table.cell(i, 1)
            cell2.text = desc

    def add_domain_info(self):
        self.document.add_page_break()
        self.document.add_heading('4. Domen haqida ma’lumot', level=1)
        
        domain_info = self.report_data.get("domain_info", {})
        
        # 4.1. Domen parametrlari
        self.document.add_heading('4.1. Domen parametrlari', level=2)
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parametr'
        hdr_cells[1].text = 'Qiymat'

        params = {
            "Domen:": self.target_domain,
            "NS server haqida ma’lumot:": "\n".join(domain_info.get("ns_servers", ["N/A"])),
            "Registrator:": domain_info.get("registrar", "N/A"),
            "Yaratilgan sana:": domain_info.get("creation_date", "N/A"),
            "Yaroqlilik muddati:": domain_info.get("expiration_date", "N/A"),
        }

        for key, value in params.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = value

        # 4.2. Saytda foydalanilayotgan texnologiyalar
        self.document.add_heading('4.2. Saytda foydalanilayotgan texnalogiyalar haqida ma’lumot', level=2)
        tech_info = self.report_data.get("tech_info", {})
        
        table_tech = self.document.add_table(rows=1, cols=2)
        table_tech.style = 'Table Grid'
        
        hdr_cells_tech = table_tech.rows[0].cells
        hdr_cells_tech[0].text = 'Komponent'
        hdr_cells_tech[1].text = 'Texnologiya'

        techs = {
            "Veb-server:": tech_info.get("server", "N/A"),
            "Dasturlash tili:": tech_info.get("lang", "N/A"),
            "Veb-dastur CMS:": tech_info.get("cms", "N/A"),
        }

        for key, value in techs.items():
            row_cells = table_tech.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = value

    def generate_report(self, filename="pentest_report.docx"):
        """Generates the full report."""
        self.add_title_page()
        self.add_table_of_contents()
        self.add_introduction()
        self.add_vulnerability_severity_levels()
        self.add_domain_info()
        # Add other sections here in the future
        
        # Save the document
        save_path = os.path.join('pentest_reports', filename)
        os.makedirs('pentest_reports', exist_ok=True)
        self.document.save(save_path)
        print(f"Report saved to {save_path}")

if __name__ == '__main__':
    # Example usage:
    target = "labka.uz"
    
    # Simulate getting domain info
    domain_details = get_whois_data(target)
    
    # This data would come from the pentesting tools
    data = {
        "vulnerabilities": [],
        "domain_info": domain_details,
        "tech_info": {
            "server": "Nginx",
            "lang": "PHP",
            "cms": "WordPress"
        }
    }
    report_gen = ReportGenerator(target, data)
    report_gen.generate_report(f"report_{target}.docx")
